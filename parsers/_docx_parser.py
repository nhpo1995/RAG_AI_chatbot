import uuid
from pathlib import Path
from typing import List
from collections import deque

from docx import Document as DocxDocument
from haystack import Document

class DocxParser:
    """
    Parser DOCX chuyên dụng: text, table, image.
    - Chunk text theo paragraph hoặc ký tự
    - Table -> Markdown với context trước
    - Image -> lưu file + context text trước
    """
    def __init__(self, images_root: Path, max_text_chunk: int = 800):
        self.images_root = images_root
        self.images_root.mkdir(exist_ok=True)
        self.max_text_chunk = max_text_chunk  # số ký tự tối đa 1 chunk

    # --------- Utils ----------
    @staticmethod
    def table_to_markdown(table) -> str:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        md = "| " + " | ".join(headers) + " |\n"
        md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for r in rows:
            md += "| " + " | ".join(r) + " |\n"
        return md.strip()

    def chunk_text(self, text: str) -> List[str]:
        """
        Chia text dài thành nhiều chunk <= max_text_chunk ký tự
        """
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.max_text_chunk
            chunks.append(text[start:end].strip())
            start = end
        return chunks

    def parse_docx(self, file_path: Path) -> List[Document]:
        docx = DocxDocument(file_path)
        file_id = str(uuid.uuid4())
        docs: List[Document] = []

        # ---------- Text ----------
        text_buffer = deque()
        for idx, para in enumerate(docx.paragraphs):
            txt = para.text.strip()
            if not txt:
                continue
            text_buffer.append(txt)

        # Chunk text từ buffer
        if text_buffer:
            combined_text = "\n\n".join(text_buffer)
            for i, chunk in enumerate(self.chunk_text(combined_text)):
                docs.append(Document(
                    content=chunk,
                    meta={
                        "category": "text",
                        "filename": file_path.name,
                        "document_id": file_id,
                        "trace": f"Text chunk {i}",
                        "source": str(file_path)
                    }
                ))

        # ---------- Table ----------
        for idx, table in enumerate(docx.tables):
            # context text trước table
            prev_text = ""
            for para in reversed(docx.paragraphs):
                if table._tbl in para._element.getparent().getchildren():
                    prev_text = para.text.strip()
                    break
            md = self.table_to_markdown(table)
            if prev_text:
                md = f"{prev_text}\n\n{md}"
            docs.append(Document(
                content=md,
                meta={
                    "category": "table",
                    "filename": file_path.name,
                    "document_id": file_id,
                    "trace": f"Table {idx}",
                    "source": str(file_path)
                }
            ))

        # ---------- Image ----------
        for idx, rel in enumerate(docx.part._rels):
            rel = docx.part._rels[rel]
            if "image" in rel.reltype:
                img_bytes = rel.target_part.blob
                images_dir = self.images_root / file_path.stem
                ext = Path(rel.target_part.partname).suffix
                img_path = images_dir / f"image_{idx}{ext}"
                img_path.parent.mkdir(parents=True, exist_ok=True)
                with open(img_path, "wb") as f:
                    f.write(img_bytes)

                # context text trước image
                prev_text = ""
                for b in reversed(docs):
                    if b.meta["category"] == "text" and b.content.strip():
                        prev_text = b.content.strip()
                        break
                content = f"Hình ảnh: {prev_text}" if prev_text else "Hình ảnh:"
                docs.append(Document(
                    content=content,
                    meta={
                        "category": "image",
                        "file_path": str(img_path),
                        "filename": file_path.name,
                        "document_id": file_id,
                        "trace": f"Hình ảnh #{idx}",
                        "source": str(file_path)
                    }
                ))

        return docs

    def run(self, folder_path: Path) -> List[Document]:
        all_docs = []
        files = [f for f in folder_path.iterdir() if f.is_file() and not f.name.startswith('.')]
        files.sort()
        for f in files:
            if f.suffix.lower() == ".docx":
                all_docs.extend(self.parse_docx(f))
        return all_docs

DATA_PATH = Path("D:/AI/haystack_data_convertor/data")  # đường dẫn chứa file DOCX
IMAGES_PATH = Path("D:/AI/haystack_data_convertor/images")  # đường dẫn lưu hình ảnh

if __name__ == "__main__":
    parser = DocxParser(images_root=IMAGES_PATH, max_text_chunk=500)
    file_path = DATA_PATH / "Đô Thị Hóa.docx"

    print(f"Parsing file: {file_path.name}\n")
    docs = parser.parse_docx(file_path)

    for idx, doc in enumerate(docs):
        print("="*40)
        print(f"{idx}. filename: {doc.meta['filename']}")
        print(f"category: {doc.meta['category']}")
        if doc.meta['category'] == "image":
            print(f"file_path: {doc.meta['file_path']}")
        print(f"trace: {doc.meta['trace']}")
        preview = doc.content if len(doc.content) < 200 else doc.content
        print(f"content preview:\n{preview}\n")
