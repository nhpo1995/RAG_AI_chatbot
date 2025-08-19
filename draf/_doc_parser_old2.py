import base64
import uuid
from collections import defaultdict
from pathlib import Path
from typing import List, Dict, Optional, Set, Any
from bs4 import BeautifulSoup
from haystack import Document
from haystack_integrations.components.converters.unstructured import (
    UnstructuredFileConverter,
)
from docx import Document as DocxDocument
from config import API_URL_UNSTRUCTURED, IMAGES_PATH, DATA_PATH


MAX_TOKENS = 500  # Số từ tối đa cho mỗi chunk


class DocParser:
    """
    Parser trích xuất text, table, image từ file (PDF, DOCX, XLSX...)
    - Text được chunk nhỏ theo số token
    - Table giữ nguyên Markdown
    - Image được trích xuất từ Word, content = đoạn text ngay trước hình ảnh
    """

    def __init__(
        self,
        api_url: str = API_URL_UNSTRUCTURED,
        images_root: Path = IMAGES_PATH,
        text_categories: Set[str] = None,
        unstructured_kwargs: Optional[Dict[str, Any]] = None,
    ):
        self.api_url = api_url
        self.images_root = images_root
        self.images_root.mkdir(exist_ok=True)
        self.text_categories = text_categories or {"Title", "NarrativeText"}
        default_unstructured_kwargs = {
            "encoding": "utf-8",
            "extract_image_block_types": ["Image", "Table"],
            "languages": ["vie", "eng"],
            "skip_infer_table_types": ["html"],
            "strategy": "hi_res",
        }
        if unstructured_kwargs:
            default_unstructured_kwargs.update(unstructured_kwargs)
        self.unstructured_kwargs = default_unstructured_kwargs
        self.converter = UnstructuredFileConverter(
            api_url=self.api_url,
            document_creation_mode="one-doc-per-element",
            separator="\n\n",
            progress_bar=False,
            unstructured_kwargs=self.unstructured_kwargs,
        )

    # ---------------- commons ----------------
    def is_text_block(self, doc: Document) -> bool:
        return doc.meta.get("category") in self.text_categories

    @staticmethod
    def chunk_text(content: str, max_tokens=MAX_TOKENS) -> List[str]:
        paragraphs = content.split("\n\n")
        chunks, chunk, token_count = [], [], 0
        for p in paragraphs:
            tokens = len(p.split())
            if token_count + tokens > max_tokens and chunk:
                chunks.append("\n\n".join(chunk))
                chunk, token_count = [], 0
            chunk.append(p)
            token_count += tokens
        if chunk:
            chunks.append("\n\n".join(chunk))
        return chunks

    def nearest_text(self, blocks: List[Document], idx: int, direction: int) -> str:
        j = idx + direction
        while 0 <= j < len(blocks):
            if self.is_text_block(blocks[j]):
                return (blocks[j].content or "").strip()
            j += direction
        return ""

    @staticmethod
    def html_table_to_markdown(html: str) -> str:
        soup = BeautifulSoup(html, "html.parser")
        table = soup.find("table")
        if not table:
            return soup.get_text(separator=" ").strip()
        rows = []
        thead = table.find("thead")
        headers = (
            [th.get_text(strip=True) for th in thead.find_all("th")] if thead else []
        )
        tbody = table.find("tbody")
        trs = tbody.find_all("tr") if tbody else table.find_all("tr")
        for i, tr in enumerate(trs):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if not headers and i == 0:
                headers = cells
                continue
            rows.append(cells)
        md_lines = []
        if headers:
            md_lines.append("| " + " | ".join(headers) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for r in rows:
            r += [""] * (len(headers) - len(r))
            md_lines.append("| " + " | ".join(r) + " |")
        return "\n".join(md_lines).strip()

    def process_table_block(
        self,
        blocks: List[Document],
        idx: int,
        file_path: Path,
        file_document_id: str,
        is_in_image: bool = False,
    ) -> Document:
        d = blocks[idx]
        md = (
            self.html_table_to_markdown(d.meta["text_as_html"])
            if "text_as_html" in d.meta
            else (d.content or "").strip()
        )
        prev_txt = self.nearest_text(blocks, idx, -1)
        if prev_txt:
            md = f"{prev_txt}\n\n{md}"
        sheet_name = d.meta.get("sheet_name")
        index = d.meta.get("element_index")
        page_num = d.meta.get("page_number")
        trace = (
            f"Table in sheet {sheet_name} with index {index}"
            if sheet_name
            else f"Table in page {page_num} with index {index}"
        )
        return Document(
            content=md,
            meta={
                "category": "table",
                "source": str(file_path),
                "filename": file_path.name,
                "document_id": file_document_id,
                "trace": trace,
                "from_image": is_in_image,
            },
        )

    @staticmethod
    def save_image(image_base64: str, out_path: Path) -> bool:
        try:
            img_bytes = base64.b64decode(image_base64)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "wb") as f:
                f.write(img_bytes)
            return True
        except Exception:
            return False

    # ---------------- Word image extraction ----------------
    def extract_images_from_docx(
        self, file_path: Path, file_document_id: str
    ) -> List[Document]:
        doc = DocxDocument(file_path)
        file_stem = file_path.stem
        images_dir = self.images_root / file_stem
        images_dir.mkdir(parents=True, exist_ok=True)
        image_docs = []
        prev_text = ""
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                prev_text = para.text.strip()
            for run in para.runs:
                if run.element.xpath(".//pic:pic"):
                    img_id = str(uuid.uuid4())
                    img_path = images_dir / f"{img_id}.png"
                    # NOTE: Lưu hình ảnh từ Word - bạn có thể dùng docx2image hoặc unstructured nếu muốn
                    # Hiện tại tạo Document với content là prev_text
                    image_docs.append(
                        Document(
                            content=prev_text,
                            meta={
                                "category": "image",
                                "file_path": str(img_path),
                                "filename": file_path.name,
                                "document_id": file_document_id,
                                "trace": f"Hình ảnh #{idx}",
                            },
                        )
                    )
        return image_docs

    # ---------------- Core ----------------
    def parse_file(self, file_path: Path) -> List[Document]:
        file_document_id = str(uuid.uuid4())
        result = self.converter.run(paths=[str(file_path)])
        blocks = result["documents"]

        # --- Text ---
        text_docs = []
        paged_texts = defaultdict(list)
        unpaged_texts = []
        for d in blocks:
            if self.is_text_block(d) and d.content:
                page_number = d.meta.get("page_number")
                if page_number is not None:
                    paged_texts[page_number].append(d.content.strip())
                else:
                    unpaged_texts.append(d.content.strip())

        # Chunk text
        for page_number, page_texts in paged_texts.items():
            merged = "\n\n".join(page_texts)
            for chunk in self.chunk_text(merged):
                text_docs.append(
                    Document(
                        content=chunk,
                        meta={
                            "category": "text",
                            "source": str(file_path),
                            "filename": file_path.name,
                            "document_id": file_document_id,
                            "trace": f"Trang {page_number}",
                        },
                    )
                )
        if unpaged_texts:
            merged = "\n\n".join(unpaged_texts)
            for chunk in self.chunk_text(merged):
                text_docs.append(
                    Document(
                        content=chunk,
                        meta={
                            "category": "text",
                            "source": str(file_path),
                            "filename": file_path.name,
                            "document_id": file_document_id,
                            "trace": "Nội dung văn bản",
                        },
                    )
                )

        # --- Tables ---
        table_docs = []
        for idx, d in enumerate(blocks):
            category = d.meta.get("category")
            if category == "Table":
                table_docs.append(
                    self.process_table_block(blocks, idx, file_path, file_document_id)
                )
            elif category == "Image" and "text_as_html" in d.meta:
                table_docs.append(
                    self.process_table_block(
                        blocks, idx, file_path, file_document_id, is_in_image=True
                    )
                )

        # --- Images (Word only) ---
        image_docs = []
        if file_path.suffix.lower() == ".docx":
            image_docs = self.extract_images_from_docx(file_path, file_document_id)

        # --- Merge ---
        final_docs = []
        final_docs.extend(text_docs)
        final_docs.extend(table_docs)
        final_docs.extend(image_docs)

        return final_docs

    def run(self, folder_path: Path) -> List[Document]:
        all_docs = []
        files = [
            f
            for f in folder_path.iterdir()
            if f.is_file() and not f.name.startswith(".")
        ]
        files.sort()
        for file_path in files:
            try:
                print(f"Parsing file: {file_path.name}")
                docs = self.parse_file(file_path)
                all_docs.extend(docs)
            except Exception as e:
                print(f"Lỗi khi xử lý file {file_path.name}: {e}")
        return all_docs


if __name__ == "__main__":
    parser = DocParser()
    file_path = DATA_PATH / "Đô Thị Hóa.docx"
    all_docs = parser.parse_file(file_path)
    for idx, doc in enumerate(all_docs):
        print(20 * "=")
        print(f"{idx}. filename= {doc.meta['filename']}")
        category = doc.meta.get("category")
        if category == "image":
            print(f"file_path: {doc.meta['file_path']}")
        print(f"category= {category}")
        print(f"trace= {doc.meta['trace']}")
        print(f"content preview: {doc.content}")
